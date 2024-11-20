from PyQt6.QtWidgets import (
    QDialog,
    QMessageBox,
    QFileDialog,
    QLineEdit,
    QPushButton,
    QVBoxLayout,
    QLabel,
    QTableView,
    QHBoxLayout,
    QSizePolicy,
    QSpacerItem,
)
from ui.py.export_window import Ui_Dialog

from docx import Document
from docx.shared import Pt, Inches
from src.base_table_model import (
    MakeModel, 
    PivotModel, 
    StatusModel, 
    RegionModel,
    GRNTIModel,
)
import os

from models import (
    select_vuz_pivot,
    select_status_pivot,
    select_region_pivot,
)


class BaseExportDialog(QDialog):
    table_model_class = MakeModel
    report = "pivot"
    report_type = 1

    def __init__(self, filter_cond={}):
        super().__init__()
        self.ui = Ui_Dialog()
        self.filters = filter_cond
        filter_trans = {
            "region": "Регион",
            "city": "Город",
            "vuz_name": "ВУЗ",
            "federation_subject": "Субъект федерации",
            "grnti_code": "Код ГРНТИ",
        }
        self.file_path = ""
        self.ui.setupUi(self)
        self.setWindowTitle("Предпросмотр отчёта")
        self.resize(1280, 720)

        # Поле для ввода имени файла
        self.file_name_line_edit = QLineEdit(self)
        self.file_name_line_edit.setPlaceholderText("Введите имя файла")

        # Кнопка выбора пути для файла
        self.browse_button = QPushButton("Выбрать путь", self)
        self.browse_button.clicked.connect(self.browse_file)

        # Текст фильтров
        self.filter_text = [QLabel(self) for _ in range(len(filter_cond))]
        for index, (key, value) in enumerate(filter_cond.items()):
            self.filter_text[index].setText(f"{filter_trans[key]}: {value}\n")

        # Поле предпросмотра
        self.preview_area = QTableView(self)
        self.table_model = self.table_model_class(
            filter_cond=filter_cond, model=self.report
        )
        self.preview_area.setModel(self.table_model)

        # Основной макет
        main_layout = QVBoxLayout()

        # Макет для выбора имени файла и кнопки выбора пути
        file_layout = QHBoxLayout()
        file_layout.addWidget(QLabel("Имя файла:"))
        file_layout.addWidget(self.file_name_line_edit)
        file_layout.addWidget(self.browse_button)

        # Лейбл и фильтрs
        filter_box = QVBoxLayout()
        filter_box.addWidget(QLabel("Фильтры:"))
        for ft in self.filter_text:
            filter_box.addWidget(ft)

        # Отступ между фильтрами и кнопками
        filter_box.addSpacerItem(
            QSpacerItem(20, 20, QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Fixed)
        )

        # Макет для кнопок установки и очистки фильтров
        filter_buttons_layout = QHBoxLayout()

        # Добавление кнопок в макет с фильтрами
        filter_box.addLayout(filter_buttons_layout)
        filter_box.addSpacerItem(
            QSpacerItem(
                20, 40, QSizePolicy.Policy.Minimum, QSizePolicy.Policy.Expanding
            )
        )

        # Макет для фильтров и предпросмотра рядом друг с другом
        preview_layout = QHBoxLayout()
        preview_layout.addLayout(filter_box)  # Добавляем фильтры слева
        preview_layout.addWidget(self.preview_area)  # Предпросмотр справа

        # Установка политики размера для предпросмотра
        self.preview_area.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding
        )

        # Макет для кнопок сохранения и отмены
        button_layout = QHBoxLayout()
        button_layout.addSpacerItem(
            QSpacerItem(
                40, 20, QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Minimum
            )
        )
        button_layout.addWidget(self.ui.pushButton_save)
        button_layout.addWidget(self.ui.pushButton_cancel)

        # Добавление всех макетов в основной
        main_layout.addLayout(file_layout)
        main_layout.addLayout(preview_layout)
        main_layout.addLayout(button_layout)

        # Применяем основной макет
        self.setLayout(main_layout)

        # Подключение кнопок к функциям
        self.ui.pushButton_save.clicked.connect(self.create_report)
        self.ui.pushButton_cancel.clicked.connect(self.confirm_close)

    def browse_file(self):
        default_file_name = self.file_name_line_edit.text()
        file_name, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить файл",
            default_file_name,
            "Документы Word (*.docx);;Все файлы (*)",
        )
        if file_name:
            # Полный путь сохраняем в атрибут
            self.file_path = file_name
            name_only = file_name.split("/")[-1]
            self.file_name_line_edit.setText(name_only)

    def confirm_close(self):
        self.close()

    def create_report(self):
        file_path = self.file_path.strip()
        if not file_path:
            self.show_notification("Пожалуйста, выберите файл и путь для сохранения.")
            return
        try:
            # Проверка доступности файла
            if os.path.exists(file_path):
                try:
                    with open(file_path, "r+"):
                        pass
                except IOError:
                    self.show_notification(
                        "Файл уже используется. Закройте его перед записью."
                    )
                    return

            if not file_path.endswith(".docx"):
                file_path += ".docx"
            print(f"Сохраняем отчёт в: {file_path}")
            make_report(
                file_path, type_report=self.report_type, filter_cond=self.filters
            )
            self.show_notification("Отчёт успешно сохранён!")
            self.close()
        except Exception as e:
            self.show_notification(f"Ошибка при сохранении отчета: {e}")

    def show_notification(self, message):
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("Уведомление")
        msg_box.setText(message)
        msg_box.setStandardButtons(QMessageBox.StandardButton.Ok)
        msg_box.setFixedHeight(500)
        msg_box.exec()


class PivotExportDialog(BaseExportDialog):
    table_model_class = PivotModel
    report = "pivot"
    report_type = 1


class StatusExportDialog(BaseExportDialog):
    table_model_class = StatusModel
    report = "status"
    report_type = 2


class RegionExportDialog(BaseExportDialog):
    table_model_class = RegionModel
    report = "region"
    report_type = 3

class GRNTIExportDialog(BaseExportDialog):
    table_model_class = GRNTIModel
    report = "grnti"
    report_type = 4


def make_report(file_path, type_report=1, filter_cond={}):
    doc = Document()
    if type_report == 1:  # 1 - По вузам
        data = select_vuz_pivot(filter_cond)
        column_names = ["Код", "ВУЗ"]
        doc.add_heading("Отчет из совдной таблицы по ВУЗам", level=1)
    elif type_report == 2:  # 2 - По статусам
        data = select_status_pivot(filter_cond)
        column_names = ["Статус"]
        doc.add_heading("Отчет из совдной таблицы по статусам", level=1)
    elif type_report == 3:  # 3 - По регионам
        data = select_region_pivot(filter_cond)
        column_names = ["Регион"]
        doc.add_heading("Отчет из совдной таблицы по регионам", level=1)
    elif type_report == 4:  # 4 - По ГРНТИ
        data = select_grnti_pivot(filter_cond)
        column_names = ["Код", "Рубрика"]
        doc.add_heading("Отчет из совдной таблицы по ГРНТИ", level=1)
    elif type_report == 5:  # 5 - По кол-ву НИР по рубрике
        data = select_vuz_pivot(filter_cond)  # должно содержать условие condrub
        column_names = ["Код", "ВУЗ"]
        doc.add_heading("Отчет из совдной таблицы по кол-ву НИР по рубрике", level=1)

    if not (filter_cond == {}):
        doc.add_heading("Фильтры:", level=2)

    if "vuz_name" in filter_cond:
        doc.add_paragraph(f"ВУЗ: {filter_cond['vuz_name']}")
    if "city" in filter_cond:
        doc.add_paragraph(f"Город: {filter_cond['city']}")
    if "federation_subject" in filter_cond:
        doc.add_paragraph(f"Субъект федерации: {filter_cond['federation_subject']}")
    if "region" in filter_cond:
        doc.add_paragraph(f"Регион: {filter_cond['region']}")

    # Настройка полей страницы
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Заголовки столбцов

    column_names.extend(
        [
            "Кол-во гр",
            "Сумма гр",
            "Кол-во НТП",
            "Сумма НТП",
            "Сумма тп",
            "Кол-во тп",
            "Общее кол-во",
            "Общая сумма",
        ]
    )

    table = doc.add_table(rows=1, cols=len(column_names))
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(column_names):
        hdr_cells[i].text = column
        for paragraph in hdr_cells[i].paragraphs:
            run = paragraph.runs[0]
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)  # размер шрифта 10
        paragraph.paragraph_format.space_after = 0

    for row in data:
        print(row.values())

    for row in data:
        row_cells = table.add_row().cells
        for i, key in enumerate(row.keys()):
            row_cells[i].text = str(row[key])

        # Установка шрифта для ячеек с данными
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    table.style = "Table Grid"
    set_table_width(table)

    doc.save(file_path)

    print("Отчет успешно создан и сохранен.")


def set_table_width(table):
    for row in table.rows:
        for cell in row.cells:
            # Устанавливаем ширину ячейки равной её контенту
            cell.width = Pt(0)  # Автоматически подстраивает по контенту
            # Убираем обтекание текста
            cell.paragraphs[0].paragraph_format.space_after = 0
