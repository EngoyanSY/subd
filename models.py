from typing import Optional, Dict
import os
import numpy as np
from docx import Document
from docx.shared import Pt

from pydantic import BaseModel
from sqlalchemy import (
    Table,
    Column,
    Integer,
    String,
    MetaData,
    create_engine,
    insert,
    select,
    func,
    union_all,
    literal_column,
)
from sqlalchemy.orm import DeclarativeBase
import pandas as pd

from schema import (
    Nir_Grant,
    Nir_NTP,
    Nir_Templan,
    Nir_VUZ,
    Nir_GRNTI,
)
from core import Session


def create_pivot(Grant, Ntp, Templan, Pivot):
    with Session() as session:
        # Первый запрос
        query1 = Grant.grant_summary()

        # Второй запрос
        query2 = Ntp.ntp_summary()

        # Третий запрос
        query3 = Templan.tp_summary()

        # Объединение запросов
        combined_query = union_all(query1, query2, query3)

        # Сводный запрос
        final_query = select(
            combined_query.c.vuz_code,
            combined_query.c.vuz_name,
            func.sum(combined_query.c.nir_grant_count).label("total_nir_grant_count"),
            func.sum(combined_query.c.total_grant_value).label("total_grant_value"),
            func.sum(combined_query.c.nir_ntp_count).label("total_nir_ntp_count"),
            func.sum(combined_query.c.total_year_value_plan).label(
                "total_year_value_plan"
            ),
            func.sum(combined_query.c.nir_templan_count).label(
                "total_nir_templan_count"
            ),
            func.sum(combined_query.c.total_value_plan).label("total_value_plan"),
            func.sum(combined_query.c.nir_grant_count)
            + func.sum(combined_query.c.nir_ntp_count)
            + func.sum(combined_query.c.nir_templan_count).label("total_count"),
            func.sum(combined_query.c.total_grant_value)
            + func.sum(combined_query.c.total_year_value_plan)
            + func.sum(combined_query.c.total_value_plan).label("total_sum"),
        ).group_by(combined_query.c.vuz_code, combined_query.c.vuz_name)

        # Выполнение запроса
        results = session.execute(final_query).fetchall()
        insert_stmt = Pivot.__table__.insert().values(
            [
                {
                    "vuz_code": row[0],
                    "vuz_name": row[1],
                    "total_nir_grant_count": row[2],
                    "total_grant_value": row[3],
                    "total_nir_ntp_count": row[4],
                    "total_year_value_plan": row[5],
                    "total_nir_templan_count": row[6],
                    "total_value_plan": row[7],
                    "total_count": row[8],
                    "total_sum": row[9],
                }
                for row in results
            ]
        )

        session.execute(insert_stmt)
        session.commit()


def set_table_width(table):
    for row in table.rows:
        for cell in row.cells:
            cell.width = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = 0


def make_report():
    with Session() as session:
        rows = session.query(Pivot).all()

        doc = Document()
        doc.add_heading("Отчет из совдной таблицы", level=1)
        table = doc.add_table(rows=1, cols=len(Pivot.__table__.columns))

        # Заголовки столбцов
        """
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(Pivot.__table__.columns):
            hdr_cells[i].text = column.name
            for paragraph in hdr_cells[i].paragraphs:
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)  # размер шрифта 10
            paragraph.paragraph_format.space_after = 0
        """

        for row in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row.UniqueID)
            row_cells[1].text = str(row.vuz_code)
            row_cells[2].text = str(row.vuz_name)
            row_cells[3].text = str(row.total_nir_grant_count)
            row_cells[4].text = str(row.total_grant_value)
            row_cells[5].text = str(row.total_nir_ntp_count)
            row_cells[6].text = str(row.total_year_value_plan)
            row_cells[7].text = str(row.total_nir_templan_count)
            row_cells[8].text = str(row.total_value_plan)
            row_cells[9].text = str(row.total_count)
            row_cells[10].text = str(row.total_sum)

            # Установка шрифта для ячеек с данными
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

        # Установка ширины столбцов
        set_table_width(table)

        doc.save("DB/report.docx")

        session.close()

    print("Отчет успешно создан и сохранен как report.docx.")


metadata_obj = MetaData()


def create_sql_tables():
    if not(os.path.isdir('DB')):
        os.mkdir("DB")
    if not (os.path.isfile("DB/DataBase.sqlite")):
        engine = create_engine("sqlite:///DB/DataBase.sqlite", echo=False)
        metadata_obj.drop_all(engine)
        metadata_obj.create_all(engine)
        gr_table = Grant()
        ntp_table = NTP()
        tp_table = Templan()
        vuz_table = VUZ()
        grnti_table = GRNTI()
        pivot_table = Pivot()
        gr_table.create_table()
        ntp_table.create_table()
        tp_table.create_table()
        grnti_table.create_table()
        vuz_table.create_table()

        create_pivot(gr_table, ntp_table, tp_table, pivot_table)
        print("База данных DataBase.sqlite создана и подключена")
    else:
        print("База данных DataBase.sqlite подключена")


class BaseTable(DeclarativeBase):
    _schema = BaseModel

    def create_table(self):
        # Считывает xlsx файл
        # Переименовывает столбцы и заменяет отсуттвующие данные на None
        # Построчно валидирует данные через Pydantic и вносит их в таблицу
        data = pd.read_excel(os.path.join("data", self._schema.table_name))
        data = data.rename(
            columns=dict(zip(data.columns, list(self._schema.model_fields)))
        )
        data.replace({np.nan: None}, inplace=True)
        with Session() as session:
            for r in range(data.shape[0]):
                table_model = self._schema.model_validate(data.iloc[r].to_dict())
                stmt = insert(self.__table__).values(table_model.dict())
                session.execute(stmt)
            session.commit()

    def select_all(self):
        return str(self.__table__.select().distinct())

    def filter(self, filter_cond: Dict):
        q = select(self.__table__)
        for fil, cond in filter_cond.items():
            if hasattr(self.__table__.c, fil):
                q = q.filter(getattr(self.__table__.c, fil).like(f"%{cond}%"))
        return q.distinct()


class Grant(BaseTable):
    _schema = Nir_Grant
    __table__ = Table(
        "nir_grant",
        metadata_obj,
        Column("UniqueID", Integer, primary_key=True),
        Column("kon_code", Integer),
        Column("nir_code", Integer),
        Column("vuz_code", Integer),
        Column("vuz_name", String),
        Column("grnti_code", String),
        Column("grant_value", Integer),
        Column("nir_director", String),
        Column("nir_name", String),
        Column("director_position", String),
        Column("director_academic_title", String),
        Column("director_academic_degree", String),
    )

    def grant_summary(self):
        return select(
            self.__table__.c.vuz_code,
            self.__table__.c.vuz_name,
            func.count(self.__table__.c.nir_code).label("nir_grant_count"),
            func.sum(self.__table__.c.grant_value).label("total_grant_value"),
            literal_column("0").label("nir_ntp_count"),
            literal_column("0").label("total_year_value_plan"),
            literal_column("0").label("nir_templan_count"),
            literal_column("0").label("total_value_plan"),
        ).group_by(self.__table__.c.vuz_code, self.__table__.c.vuz_name)

    def get_grnti_codes(self, filter_cond: Optional[Dict] = None):
        q = select(self.__table__.c.grnti_code).join(
            VUZ, VUZ.vuz_code == self.__table__.c.vuz_code, isouter=True
        )
        if filter_cond is not None:
            q = q.filter_by(**filter_cond)
        return q.order_by(self.__table__.c.grnti_code).distinct()


class NTP(BaseTable):
    _schema = Nir_NTP
    __table__ = Table(
        "nir_ntp",
        metadata_obj,
        Column("UniqueID", Integer, primary_key=True),
        Column("ntp_code", Integer),
        Column("nir_number", Integer),
        Column("vuz_code", Integer),
        Column("vuz_name", String),
        Column("grnti_code", String),
        Column("year_value_plan", Integer),
        Column("nir_director", String),
        Column("nir_type", String),
        Column("nir_name", String),
        Column("director_meta", String),
    )

    def ntp_summary(self):
        return select(
            self.__table__.c.vuz_code,
            self.__table__.c.vuz_name,
            literal_column("0").label("nir_grant_count"),
            literal_column("0").label("total_grant_value"),
            func.count(self.__table__.c.nir_number).label("nir_ntp_count"),
            func.sum(self.__table__.c.year_value_plan).label("total_year_value_plan"),
            literal_column("0").label("nir_templan_count"),
            literal_column("0").label("total_value_plan"),
        ).group_by(self.__table__.c.vuz_code, self.__table__.c.vuz_name)


class Templan(BaseTable):
    _schema = Nir_Templan
    __table__ = Table(
        "nir_templan",
        metadata_obj,
        Column("UniqueID", Integer, primary_key=True),
        Column("vuz_code", Integer),
        Column("vuz_name", String),
        Column("grnti_code", String),
        Column("value_plan", Integer),
        Column("nir_director", String),
        Column("nir_type", String),
        Column("nir_reg_number", String),
        Column("nir_name", String),
        Column("director_position", String),
    )

    def tp_summary(self):
        return select(
            self.__table__.c.vuz_code,
            self.__table__.c.vuz_name,
            literal_column("0").label("nir_grant_count"),
            literal_column("0").label("total_grant_value"),
            literal_column("0").label("nir_ntp_count"),
            literal_column("0").label("total_year_value_plan"),
            func.count(self.__table__.c.nir_reg_number).label("nir_templan_count"),
            func.sum(self.__table__.c.value_plan).label("total_value_plan"),
        ).group_by(self.__table__.c.vuz_code, self.__table__.c.vuz_name)


class VUZ(BaseTable):
    _schema = Nir_VUZ
    __table__ = Table(
        "vuz",
        metadata_obj,
        Column("UniqueID", Integer, primary_key=True),
        Column("vuz_code", Integer),
        Column("vuz_name", String),
        Column("status", String),
        Column("fed_sub_code", Integer),
        Column("federation_subject", String),
        Column("region", String),
        Column("city", String),
        Column("name", String),
        Column("full_name", String),
        Column("gr_ved", String),
        Column("profile", String),
    )


class GRNTI(BaseTable):
    _schema = Nir_GRNTI
    __table__ = Table(
        "grnti",
        metadata_obj,
        Column("UniqueID", Integer, primary_key=True),
        Column("codrub", String),
        Column("rubrika", String),
    )


class Pivot(BaseTable):
    __table__ = Table(
        "pivot",
        metadata_obj,
        Column("UniqueID", Integer, primary_key=True),
        Column("vuz_code", Integer),
        Column("vuz_name", String),
        Column("total_nir_grant_count", Integer),
        Column("total_grant_value", Integer),
        Column("total_nir_ntp_count", Integer),
        Column("total_year_value_plan", Integer),
        Column("total_nir_templan_count", Integer),
        Column("total_value_plan", Integer),
        Column("total_count", Integer),
        Column("total_sum", Integer),
    )
